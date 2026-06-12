import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_md_to_docx(md_path, docx_path):
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return

    doc = Document()
    
    # Set Margins (4cm left, 3cm others)
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1.57) # ~4cm
        section.right_margin = Inches(1.18) # ~3cm
        section.top_margin = Inches(1.18) # ~3cm
        section.bottom_margin = Inches(1.18) # ~3cm

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.bold = True
                run.font.color.rgb = None # Black
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=1)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = None
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=2)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = None
        
        # Horizontal Rule
        elif line == '---':
            doc.add_page_break()
            
        # List items
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
        elif re.match(r'^\d+\.', line):
            p = doc.add_paragraph(line[line.find('.')+2:], style='List Number')
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        # Normal text
        else:
            # Clean bold/italic markers for simple conversion
            clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
            clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)
            
            p = doc.add_paragraph(clean_line)
            p.paragraph_format.line_spacing = 1.15
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    doc.save(docx_path)
    print(f"Successfully converted {md_path} to {docx_path}")

if __name__ == "__main__":
    convert_md_to_docx('Laporan.md', 'Laporan.docx')
